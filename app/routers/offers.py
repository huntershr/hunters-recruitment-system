from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import StreamingResponse
from sqlalchemy.orm import Session, joinedload
from typing import Optional
from pydantic import BaseModel
from datetime import datetime
import io
import logging

from .. import models, database
from ..routers.auth import get_current_user

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/api/admin/offers", tags=["Offers"])


def get_db():
    db = database.SessionLocal()
    try:
        yield db
    finally:
        db.close()


class OfferIn(BaseModel):
    application_id: int
    candidate_name: str
    job_title: str
    department: Optional[str] = ""
    start_date: Optional[str] = ""
    working_hours_from: Optional[str] = "9:00"
    working_hours_to: Optional[str] = "5:00"
    net_salary: Optional[str] = ""
    reporting_to: Optional[str] = ""
    exceptions: Optional[str] = ""


class OfferStatusIn(BaseModel):
    status: str  # accepted / rejected / pending


def _offer_to_dict(offer) -> dict:
    return {
        "id": offer.id,
        "application_id": offer.application_id,
        "candidate_name": offer.candidate_name,
        "job_title": offer.job_title,
        "department": offer.department,
        "start_date": offer.start_date,
        "working_hours_from": offer.working_hours_from,
        "working_hours_to": offer.working_hours_to,
        "net_salary": offer.net_salary,
        "reporting_to": offer.reporting_to,
        "exceptions": offer.exceptions,
        "status": offer.status,
        "created_at": offer.created_at.isoformat() if offer.created_at else None,
        "created_by": offer.created_by,
    }


def _check_scope(application_id: int, current_user, db: Session) -> bool:
    if current_user.is_admin:
        return True
    if not current_user.company_id:
        return False
    app = db.query(models.Application).filter(models.Application.id == application_id).first()
    if not app or not app.job:
        return False
    owner = db.query(models.User).filter(models.User.id == app.job.owner_id).first()
    return bool(owner and owner.company_id == current_user.company_id)


def _hex_to_rgb(h):
    h = h.lstrip('#')
    return tuple(int(h[i:i+2], 16) for i in (0, 2, 4))


def _set_cell_bg(cell, hex_color):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for shd in tcPr.findall(qn('w:shd')):
        tcPr.remove(shd)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color.lstrip('#'))
    tcPr.append(shd)


def _set_cell_borders(cell, color='1B2A4A', size='4'):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top', 'left', 'bottom', 'right']:
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), size)
        b.set(qn('w:color'), color)
        tcBorders.append(b)
    tcPr.append(tcBorders)


def _set_no_borders(cell):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top', 'left', 'bottom', 'right']:
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'none')
        b.set(qn('w:sz'), '0')
        b.set(qn('w:color'), 'FFFFFF')
        tcBorders.append(b)
    tcPr.append(tcBorders)


def _add_para_border(para, side, color, size='12'):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    b = OxmlElement(f'w:{side}')
    b.set(qn('w:val'), 'single')
    b.set(qn('w:sz'), size)
    b.set(qn('w:color'), color)
    b.set(qn('w:space'), '1')
    pBdr.append(b)
    pPr.append(pBdr)


def _cell_para(cell, text, bold=False, color='000000', size=9.5, align=None):
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    para = cell.paragraphs[0]
    para.clear()
    if align is not None:
        para.alignment = align
    run = para.add_run(text)
    run.font.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(*_hex_to_rgb(color))
    run.font.name = 'Calibri'
    return para


def generate_offer_docx(offer, candidate_name, job_title, company_name, company_logo_url=None):
    import requests as req_lib
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()

    NAVY = '1B2A4A'
    GOLD = 'C9A84C'
    WHITE = 'FFFFFF'
    LIGHT = 'F0F4F8'

    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    # ── HEADER TABLE (logo left | company name right) ──
    hdr = doc.add_table(rows=1, cols=2)
    hdr.autofit = False
    hdr.columns[0].width = Cm(5.5)
    hdr.columns[1].width = Cm(12.0)

    logo_cell = hdr.cell(0, 0)
    _set_cell_bg(logo_cell, NAVY)
    _set_no_borders(logo_cell)
    # Gold right border as divider between logo and company name
    tcPr = logo_cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    right_b = OxmlElement('w:right')
    right_b.set(qn('w:val'), 'single')
    right_b.set(qn('w:sz'), '4')
    right_b.set(qn('w:color'), GOLD)
    tcBorders.append(right_b)
    tcPr.append(tcBorders)

    logo_para = logo_cell.paragraphs[0]
    logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    logo_para.paragraph_format.space_before = Pt(8)
    logo_para.paragraph_format.space_after = Pt(8)
    if company_logo_url:
        try:
            r = req_lib.get(company_logo_url, timeout=5)
            img_stream = io.BytesIO(r.content)
            logo_para.add_run().add_picture(img_stream, width=Cm(3.8))
        except Exception:
            run = logo_para.add_run(company_name or '')
            run.font.color.rgb = RGBColor(*_hex_to_rgb(WHITE))
            run.font.size = Pt(11)
            run.font.bold = True
            run.font.name = 'Calibri'
    else:
        run = logo_para.add_run(company_name or '')
        run.font.color.rgb = RGBColor(*_hex_to_rgb(WHITE))
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.name = 'Calibri'

    info_cell = hdr.cell(0, 1)
    _set_cell_bg(info_cell, NAVY)
    _set_no_borders(info_cell)
    info_para = info_cell.paragraphs[0]
    info_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    info_para.paragraph_format.space_before = Pt(10)
    run = info_para.add_run(company_name or '')
    run.font.color.rgb = RGBColor(*_hex_to_rgb(WHITE))
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.name = 'Calibri'

    sub_para = info_cell.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    sub_para.paragraph_format.space_after = Pt(10)
    sub_run = sub_para.add_run('Job Offer Letter')
    sub_run.font.color.rgb = RGBColor(*_hex_to_rgb(GOLD))
    sub_run.font.size = Pt(10)
    sub_run.font.italic = True
    sub_run.font.name = 'Calibri'

    # ── GOLD DIVIDER ──
    div = doc.add_paragraph()
    div.paragraph_format.space_before = Pt(0)
    div.paragraph_format.space_after = Pt(0)
    _add_para_border(div, 'bottom', GOLD, '12')

    # ── DATE ──
    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    date_p.paragraph_format.space_before = Pt(6)
    date_p.paragraph_format.space_after = Pt(4)
    dr = date_p.add_run(f'Date: {offer.created_at.strftime("%B %d, %Y") if offer.created_at else ""}')
    dr.font.size = Pt(8.5)
    dr.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    dr.font.name = 'Calibri'

    # ── GREETING ──
    greet = doc.add_paragraph()
    greet.paragraph_format.space_before = Pt(10)
    greet.paragraph_format.space_after = Pt(6)
    gr = greet.add_run(f'Dear Ms / Mr {candidate_name},')
    gr.font.size = Pt(11)
    gr.font.bold = True
    gr.font.color.rgb = RGBColor(*_hex_to_rgb(NAVY))
    gr.font.name = 'Calibri'

    # ── INTRO ──
    for text in (
        f'It is with great pleasure that we are offering you the position of '
        f'{job_title} with {company_name}. '
        f'Your experience and enthusiasm will be a valuable asset to our organization.',
        'Please review the offer details below, sign where indicated, and return a copy to HR.',
    ):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(5)
        r = p.add_run(text)
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        r.font.name = 'Calibri'

    # ── OFFER TABLE ──
    rows_data = [
        ('Candidate Name', candidate_name or ''),
        ('Job Title', job_title or ''),
        ('Department / Division / Site', offer.department or ''),
        ('Proposed Starting Date', str(offer.start_date) if offer.start_date else ''),
        ('Working Days', 'Sunday to Thursday'),
        ('Working Hours', f'{offer.working_hours_from or ""} to {offer.working_hours_to or ""}'),
        ('Package Details / Net Salary', f'{offer.net_salary or ""} EGP — Net (monthly by direct transfer to payroll bank account)'),
        ('Exceptions and Permissions', offer.exceptions or 'None'),
        ('Reporting To', offer.reporting_to or ''),
    ]

    tbl = doc.add_table(rows=len(rows_data) + 1, cols=2)
    tbl.autofit = False
    tbl.columns[0].width = Cm(6.5)
    tbl.columns[1].width = Cm(11.0)

    for i, hdr_text in enumerate(['Field', 'Details']):
        c = tbl.rows[0].cells[i]
        _set_cell_bg(c, NAVY)
        _set_cell_borders(c, NAVY)
        _cell_para(c, hdr_text, bold=True, color=WHITE, size=9.5,
                   align=WD_ALIGN_PARAGRAPH.LEFT)

    for idx, (label, value) in enumerate(rows_data):
        bg = LIGHT if idx % 2 == 0 else WHITE
        row = tbl.rows[idx + 1]
        lc, vc = row.cells[0], row.cells[1]
        _set_cell_bg(lc, bg)
        _set_cell_bg(vc, bg)
        _set_cell_borders(lc, 'CCCCCC', '2')
        _set_cell_borders(vc, 'CCCCCC', '2')
        _cell_para(lc, label, bold=True, color=NAVY, size=9.5)
        _cell_para(vc, value, bold=False, color='333333', size=9.5)

    # ── REQUIRED DOCUMENTS ──
    doc.add_paragraph()
    docs_title = doc.add_paragraph()
    docs_title.paragraph_format.space_before = Pt(8)
    docs_title.paragraph_format.space_after = Pt(4)
    # Gold left-border accent
    pPr = docs_title._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    left_b = OxmlElement('w:left')
    left_b.set(qn('w:val'), 'single')
    left_b.set(qn('w:sz'), '18')
    left_b.set(qn('w:color'), GOLD)
    left_b.set(qn('w:space'), '6')
    pBdr.append(left_b)
    pPr.append(pBdr)
    dtr = docs_title.add_run('Required Hiring Documents:')
    dtr.font.bold = True
    dtr.font.size = Pt(10)
    dtr.font.color.rgb = RGBColor(*_hex_to_rgb(NAVY))
    dtr.font.name = 'Calibri'

    arabic_docs = [
        'أصل شهادة الميلاد',
        'أصل شهادة المؤهل الدراسي',
        'أصل شهادة الخدمة العسكرية للذكور',
        'كعب عمل',
        'صورة بطاقة الرقم القومي',
        'عدد 6 صور شخصية حديثة',
        'صحيفة الحالة الجنائية',
        'بيان بآخر راتب / مفردات مرتب',
        'نموذج 111 خاص بالتأمين الصحي',
    ]
    for item in arabic_docs:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(item)
        r.font.size = Pt(9.5)
        r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        r.font.name = 'Calibri'

    # ── GOLD DIVIDER BEFORE SIGNATURES ──
    sig_div = doc.add_paragraph()
    sig_div.paragraph_format.space_before = Pt(14)
    sig_div.paragraph_format.space_after = Pt(8)
    _add_para_border(sig_div, 'top', GOLD, '8')

    # ── SIGNATURES TABLE ──
    sig_tbl = doc.add_table(rows=1, cols=2)
    sig_tbl.autofit = False
    sig_tbl.columns[0].width = Cm(8.5)
    sig_tbl.columns[1].width = Cm(9.0)

    for col_idx, title in enumerate(['Candidate Acceptance', 'Offered By']):
        sc = sig_tbl.cell(0, col_idx)
        _set_no_borders(sc)
        tp = sc.paragraphs[0]
        tr = tp.add_run(title.upper())
        tr.font.bold = True
        tr.font.size = Pt(9)
        tr.font.color.rgb = RGBColor(*_hex_to_rgb(NAVY))
        tr.font.name = 'Calibri'
        for line_label in ['Full Name', 'Signature', 'Date']:
            lp = sc.add_paragraph()
            lp.paragraph_format.space_before = Pt(10)
            llr = lp.add_run(f'{line_label}:  ')
            llr.font.size = Pt(8.5)
            llr.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
            llr.font.name = 'Calibri'
            blank = lp.add_run('_' * 28)
            blank.font.size = Pt(8.5)
            blank.font.color.rgb = RGBColor(0xBB, 0xBB, 0xBB)

    # ── FOOTER TABLE (navy background) ──
    footer_tbl = doc.add_table(rows=1, cols=2)
    footer_tbl.autofit = False
    footer_tbl.columns[0].width = Cm(9.5)
    footer_tbl.columns[1].width = Cm(8.0)

    lfc = footer_tbl.cell(0, 0)
    rfc = footer_tbl.cell(0, 1)
    for fc in (lfc, rfc):
        _set_cell_bg(fc, NAVY)
        _set_no_borders(fc)

    lfp = lfc.paragraphs[0]
    lfp.paragraph_format.space_before = Pt(6)
    lfp.paragraph_format.space_after = Pt(6)
    lfr = lfp.add_run('Powered by Hunters HR  |  hunters-egypt.com')
    lfr.font.size = Pt(7.5)
    lfr.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
    lfr.font.italic = True
    lfr.font.name = 'Calibri'

    rfp = rfc.paragraphs[0]
    rfp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rfp.paragraph_format.space_before = Pt(6)
    rfp.paragraph_format.space_after = Pt(6)
    rfr = rfp.add_run('hr@hunters-egypt.com')
    rfr.font.size = Pt(7.5)
    rfr.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
    rfr.font.italic = True
    rfr.font.name = 'Calibri'

    return doc


# ── Routes ──────────────────────────────────────────────────────────────────

@router.post("/")
def create_offer(
    data: OfferIn,
    db: Session = Depends(get_db),
    current_user: models.User = Depends(get_current_user),
):
    if not current_user.is_admin and not current_user.company_id:
        raise HTTPException(status_code=403, detail="Access denied")
    if not _check_scope(data.application_id, current_user, db):
        raise HTTPException(status_code=403, detail="Access denied")

    existing = db.query(models.Offer).filter(models.Offer.application_id == data.application_id).first()
    if existing:
        for field in ["candidate_name", "job_title", "department", "start_date",
                      "working_hours_from", "working_hours_to", "net_salary",
                      "reporting_to", "exceptions"]:
            setattr(existing, field, getattr(data, field))
        db.commit()
        db.refresh(existing)
        logger.info(f"Offer {existing.id} updated for application {data.application_id}")
        return _offer_to_dict(existing)

    offer = models.Offer(
        application_id=data.application_id,
        candidate_name=data.candidate_name,
        job_title=data.job_title,
        department=data.department,
        start_date=data.start_date,
        working_hours_from=data.working_hours_from,
        working_hours_to=data.working_hours_to,
        net_salary=data.net_salary,
        reporting_to=data.reporting_to,
        exceptions=data.exceptions,
        status="pending",
        created_by=current_user.id,
    )
    db.add(offer)
    db.commit()
    db.refresh(offer)
    logger.info(f"Offer {offer.id} created for application {data.application_id}")
    return _offer_to_dict(offer)


@router.get("/application/{application_id}")
def get_offer_by_application(
    application_id: int,
    db: Session = Depends(get_db),
    current_user: models.User = Depends(get_current_user),
):
    if not current_user.is_admin and not current_user.company_id:
        raise HTTPException(status_code=403, detail="Access denied")
    offer = db.query(models.Offer).filter(models.Offer.application_id == application_id).first()
    if not offer:
        return None
    return _offer_to_dict(offer)


@router.patch("/{offer_id}/status")
def update_offer_status(
    offer_id: int,
    data: OfferStatusIn,
    db: Session = Depends(get_db),
    current_user: models.User = Depends(get_current_user),
):
    if not current_user.is_admin and not current_user.company_id:
        raise HTTPException(status_code=403, detail="Access denied")
    offer = db.query(models.Offer).filter(models.Offer.id == offer_id).first()
    if not offer:
        raise HTTPException(status_code=404, detail="Offer not found")
    if data.status not in ("accepted", "rejected", "pending"):
        raise HTTPException(status_code=400, detail="Invalid status — use accepted / rejected / pending")
    offer.status = data.status
    db.commit()
    db.refresh(offer)
    logger.info(f"Offer {offer_id} status → {data.status}")
    return _offer_to_dict(offer)


@router.get("/{offer_id}/download")
def download_offer(
    offer_id: int,
    db: Session = Depends(get_db),
    current_user: models.User = Depends(get_current_user),
):
    if not current_user.is_admin and not current_user.company_id:
        raise HTTPException(status_code=403, detail="Access denied")
    offer = db.query(models.Offer).filter(models.Offer.id == offer_id).first()
    if not offer:
        raise HTTPException(status_code=404, detail="Offer not found")

    company_name = "Hunters for HR Transformation & Execution"
    logo_url = None
    app = (
        db.query(models.Application)
        .options(
            joinedload(models.Application.job)
                .joinedload(models.Job.owner)
                .joinedload(models.User.company)
        )
        .filter(models.Application.id == offer.application_id)
        .first()
    )
    if app and app.job and app.job.owner and app.job.owner.company:
        company = app.job.owner.company
        company_name = company.company_name
        logo_url = company.logo_url if company else None

    doc = generate_offer_docx(offer, offer.candidate_name, offer.job_title, company_name, logo_url)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    safe_name = (offer.candidate_name or "Offer").replace(" ", "_")
    filename = f"JobOffer_{safe_name}.docx"

    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )
