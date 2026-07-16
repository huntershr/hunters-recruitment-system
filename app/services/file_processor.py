import pandas as pd
from pypdf import PdfReader
import docx
import io
import re
import logging

logger = logging.getLogger(__name__)

# Suffix must be 1-2 all-consonant chars (y excluded, treated as vowel).
# Catches "Suppo rt", "Learne r", "Managemen t" while ignoring common English
# prepositions/articles whose suffixes contain vowels ("Expert in", "Due to", etc.).
_MIDWORD_SPLIT_RE = re.compile(r'\b[A-Za-z]{4,} [bcdfghjklmnpqrstvwxz]{1,2}\b')

_TYPE_A_THRESHOLD = 0.08   # Latin single-alpha ratio
_TYPE_B_THRESHOLD = 0.03   # Mid-word consonant-only fragment ratio


def _is_chromium_pdf(reader):
    """Metadata fast-path: True for Chromium/Skia/Canva-rendered PDFs."""
    try:
        meta = reader.metadata or {}
        producer = str(meta.get("/Producer") or "").lower()
        creator  = str(meta.get("/Creator")  or "").lower()
        for kw in ("chromium", "skia", "chrome", "canva"):
            if kw in producer or kw in creator:
                return True
    except Exception:
        pass
    return False


def _symptom_score(text):
    """
    Return (type_a_ratio, type_b_ratio, description).

    Type A: fraction of Latin-script tokens that are single alphabetic characters.
            High ratios indicate char-per-text-object rendering (Chromium/Canva style).
            Latin-script guard prevents false positives on Arabic-language CVs.
    Type B: fraction of tokens preceded by a word fragment ending in 1-2 consonants
            with no vowels (e.g. "Suppo rt", "Learne r").
            Indicates arbitrary-boundary splits from Nitro Pro / Adobe-edited PDFs.
    """
    if not text:
        return 0.0, 0.0, "empty"
    sample = text[:8000]
    tokens = sample.split()
    if not tokens:
        return 0.0, 0.0, "no tokens"

    latin_single = sum(1 for t in tokens if len(t) == 1 and t.isalpha() and ord(t) < 128)
    latin_total  = sum(1 for t in tokens if any(ord(c) < 128 and c.isalpha() for c in t))
    type_a = latin_single / max(1, latin_total)

    mid_word_hits = len(_MIDWORD_SPLIT_RE.findall(sample))
    type_b = mid_word_hits / max(1, len(tokens))

    desc = (
        f"type_a={type_a:.1%}({latin_single}/{latin_total} Latin tokens), "
        f"type_b={type_b:.1%}({mid_word_hits}/{len(tokens)} tokens)"
    )
    return type_a, type_b, desc


def _extract_all_pages(reader, use_layout):
    text = ""
    for page in reader.pages:
        try:
            page_text = (
                page.extract_text(extraction_mode="layout")
                if use_layout
                else page.extract_text()
            )
        except TypeError:
            page_text = page.extract_text()
        if page_text:
            text += page_text + "\n"
    return text.strip()


def extract_text_from_pdf(file_content):
    try:
        reader = PdfReader(io.BytesIO(file_content))

        # Fast-path: Chromium/Canva metadata confirmed → default mode is already correct
        if _is_chromium_pdf(reader):
            text = _extract_all_pages(reader, use_layout=False)
            logger.info("PDF extraction: default mode (Chromium/Canva metadata fast-path)")
            return text

        # Primary pass: default mode
        text_default = _extract_all_pages(reader, use_layout=False)
        ta_def, tb_def, reason_def = _symptom_score(text_default)

        if ta_def < _TYPE_A_THRESHOLD and tb_def < _TYPE_B_THRESHOLD:
            logger.info(f"PDF extraction: default mode (no symptoms). {reason_def}")
            return text_default

        # Symptoms detected → also try layout mode and pick lower-symptom result
        text_layout = _extract_all_pages(reader, use_layout=True)
        ta_lay, tb_lay, reason_lay = _symptom_score(text_layout)

        if (ta_lay + tb_lay) <= (ta_def + tb_def):
            logger.info(
                f"PDF extraction: layout mode chosen (symptoms in default). "
                f"default=[{reason_def}] layout=[{reason_lay}]"
            )
            return text_layout
        else:
            logger.info(
                f"PDF extraction: default mode retained (layout mode worse). "
                f"default=[{reason_def}] layout=[{reason_lay}]"
            )
            return text_default

    except Exception as e:
        logger.error(f"Error extracting PDF: {e}")
        return ""

def extract_text_from_docx(file_content):
    try:
        doc = docx.Document(io.BytesIO(file_content))
        parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)
        # Also extract text from tables (two-column DOCX CVs use table layouts)
        for table in doc.tables:
            for row in table.rows:
                row_parts = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_parts:
                    parts.append("  ".join(row_parts))
        return "\n".join(parts).strip()
    except Exception as e:
        logger.error(f"Error extracting Word: {e}")
        return ""

def process_excel_candidates(file_content):
    """
    Returns a list of dictionaries from Excel for candidates.
    Expected columns: Name, Email, Phone, Experience, Education, Skills, Salary
    """
    try:
        df = pd.read_excel(io.BytesIO(file_content))
        # Convert NaN to empty string
        df = df.fillna("")
        return df.to_dict(orient="records")
    except Exception as e:
        logger.error(f"Error processing Excel candidates: {e}")
        return []

def process_excel_jobs(file_content):
    """
    Returns a list of dictionaries from Excel for jobs.
    Expected columns: Title, Description, Skills, Experience, Education, Salary Range
    """
    try:
        df = pd.read_excel(io.BytesIO(file_content))
        df = df.fillna("")
        return df.to_dict(orient="records")
    except Exception as e:
        logger.error(f"Error processing Excel jobs: {e}")
        return []

def extract_text_from_file(filename, file_content):
    if filename.endswith(".pdf"):
        return extract_text_from_pdf(file_content)
    elif filename.endswith(".docx") or filename.endswith(".doc"):
        return extract_text_from_docx(file_content)
    return ""
